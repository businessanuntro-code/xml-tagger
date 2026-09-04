from fastapi import FastAPI, File, UploadFile, Request
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from docx import Document
from docx.oxml.ns import qn

from pathlib import Path
from html import escape
import shutil
import uuid
import re


# ============================================================
# CONFIG
# ============================================================

BASE_DIR = Path(__file__).resolve().parent

TEMPLATES_DIR = BASE_DIR / "templates"
UPLOADS_DIR = BASE_DIR / "uploads"

WORD_UPLOADS_DIR = UPLOADS_DIR / "word"
WORD_MEDIA_DIR = UPLOADS_DIR / "word_media"

WORD_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
WORD_MEDIA_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================
# FASTAPI
# ============================================================

app = FastAPI(
    title="XML Tagger",
    version="3.0.0"
)

templates = Jinja2Templates(directory=str(TEMPLATES_DIR))


# ============================================================
# STATIC FILES
# ============================================================

app.mount(
    "/files/word",
    StaticFiles(directory=str(WORD_UPLOADS_DIR)),
    name="word_files"
)

app.mount(
    "/files/media",
    StaticFiles(directory=str(WORD_MEDIA_DIR)),
    name="word_media"
)


# ============================================================
# HELPERS
# ============================================================

def safe_filename(filename: str) -> str:
    """
    Curata numele fisierului pentru a evita caractere problematice.
    """

    if not filename:
        return "document.docx"

    filename = Path(filename).name

    filename = re.sub(
        r"[^a-zA-Z0-9_\-.ăâîșțĂÂÎȘȚ ]+",
        "_",
        filename
    )

    filename = filename.replace(" ", "_")

    return filename


def escape_html(value) -> str:
    if value is None:
        return ""

    return escape(str(value))


def twips_to_px(value):
    """
    Word foloseste twips.
    Aproximam 96 DPI.
    1 inch = 1440 twips
    1 inch = 96 px
    => 1 twip = 0.066666 px
    """

    if value is None:
        return 0

    try:
        return round(float(value) * 96 / 1440, 2)
    except Exception:
        return 0


def half_points_to_px(value):
    """
    Word font size este in half-points.
    12 pt => 24 half-points.
    """

    if value is None:
        return None

    try:
        points = float(value) / 2
        return round(points * 96 / 72, 2)
    except Exception:
        return None


def color_to_css(color):
    if not color:
        return None

    try:
        rgb = color.rgb

        if rgb:
            return "#" + str(rgb)

    except Exception:
        pass

    return None


# ============================================================
# RUN FORMATTING
# ============================================================

def run_to_html(run):
    """
    Transforma un run Word in HTML pastrand cat mai mult
    din formatarea originala.
    """

    text = run.text or ""

    if not text:
        return ""

    text = escape_html(text)

    styles = []

    # --------------------------------------------------------
    # FONT
    # --------------------------------------------------------

    try:
        font_name = run.font.name

        if font_name:
            styles.append(
                f"font-family:'{escape_html(font_name)}';"
            )
    except Exception:
        pass

    # --------------------------------------------------------
    # FONT SIZE
    # --------------------------------------------------------

    try:
        if run.font.size:
            px = half_points_to_px(run.font.size.pt * 2)

            if px:
                styles.append(f"font-size:{px}px;")
    except Exception:
        pass

    # --------------------------------------------------------
    # COLOR
    # --------------------------------------------------------

    try:
        css_color = color_to_css(run.font.color)

        if css_color:
            styles.append(f"color:{css_color};")
    except Exception:
        pass

    # --------------------------------------------------------
    # BOLD
    # --------------------------------------------------------

    try:
        if run.bold is True:
            styles.append("font-weight:700;")
        elif run.bold is False:
            styles.append("font-weight:400;")
    except Exception:
        pass

    # --------------------------------------------------------
    # ITALIC
    # --------------------------------------------------------

    try:
        if run.italic is True:
            styles.append("font-style:italic;")
        elif run.italic is False:
            styles.append("font-style:normal;")
    except Exception:
        pass

    # --------------------------------------------------------
    # UNDERLINE
    # --------------------------------------------------------

    try:
        if run.underline:
            styles.append("text-decoration:underline;")
    except Exception:
        pass

    # --------------------------------------------------------
    # STRIKE
    # --------------------------------------------------------

    try:
        if run.font.strike:
            styles.append("text-decoration:line-through;")
    except Exception:
        pass

    # --------------------------------------------------------
    # SUPERSCRIPT
    # --------------------------------------------------------

    try:
        if run.font.superscript:
            styles.append(
                "vertical-align:super;"
                "font-size:0.75em;"
            )
    except Exception:
        pass

    # --------------------------------------------------------
    # SUBSCRIPT
    # --------------------------------------------------------

    try:
        if run.font.subscript:
            styles.append(
                "vertical-align:sub;"
                "font-size:0.75em;"
            )
    except Exception:
        pass

    # --------------------------------------------------------
    # SMALL CAPS
    # --------------------------------------------------------

    try:
        if run.font.small_caps:
            styles.append("font-variant:small-caps;")
    except Exception:
        pass

    # --------------------------------------------------------
    # HIGHLIGHT
    # --------------------------------------------------------

    try:
        highlight = run.font.highlight_color

        if highlight:
            highlight_name = str(highlight).lower()

            highlight_map = {
                "yellow": "#ffff00",
                "green": "#00ff00",
                "cyan": "#00ffff",
                "magenta": "#ff00ff",
                "blue": "#0000ff",
                "red": "#ff0000",
                "dark_blue": "#000080",
                "dark_red": "#800000",
                "dark_green": "#008000",
                "dark_yellow": "#808000",
                "dark_cyan": "#008080",
                "dark_magenta": "#800080",
                "gray_25": "#c0c0c0",
                "gray_50": "#808080",
            }

            css = highlight_map.get(highlight_name)

            if css:
                styles.append(
                    f"background-color:{css};"
                )
    except Exception:
        pass

    # --------------------------------------------------------
    # CHARACTER SPACING
    # --------------------------------------------------------

    style_attr = "".join(styles)

    if style_attr:
        return f'<span style="{style_attr}">{text}</span>'

    return text


# ============================================================
# PARAGRAPH ALIGNMENT
# ============================================================

def get_alignment(paragraph):
    try:
        alignment = paragraph.alignment

        if alignment is None:
            return "left"

        value = str(alignment).lower()

        if "center" in value:
            return "center"

        if "right" in value:
            return "right"

        if "justify" in value:
            return "justify"

        if "distribute" in value:
            return "justify"

    except Exception:
        pass

    return "left"


# ============================================================
# PARAGRAPH STYLE
# ============================================================

def get_paragraph_css(paragraph):
    styles = []

    # --------------------------------------------------------
    # ALIGNMENT
    # --------------------------------------------------------

    styles.append(
        f"text-align:{get_alignment(paragraph)};"
    )

    # --------------------------------------------------------
    # INDENTATION
    # --------------------------------------------------------

    try:
        fmt = paragraph.paragraph_format

        if fmt.left_indent:
            px = twips_to_px(fmt.left_indent.twips)

            if px:
                styles.append(
                    f"margin-left:{px}px;"
                )

        if fmt.right_indent:
            px = twips_to_px(fmt.right_indent.twips)

            if px:
                styles.append(
                    f"margin-right:{px}px;"
                )

        if fmt.first_line_indent:
            px = twips_to_px(fmt.first_line_indent.twips)

            if px:
                styles.append(
                    f"text-indent:{px}px;"
                )
    except Exception:
        pass

    # --------------------------------------------------------
    # SPACING BEFORE / AFTER
    # --------------------------------------------------------

    try:
        fmt = paragraph.paragraph_format

        if fmt.space_before:
            px = twips_to_px(fmt.space_before.twips)

            styles.append(
                f"margin-top:{px}px;"
            )

        if fmt.space_after:
            px = twips_to_px(fmt.space_after.twips)

            styles.append(
                f"margin-bottom:{px}px;"
            )
    except Exception:
        pass

    # --------------------------------------------------------
    # LINE SPACING
    # --------------------------------------------------------

    try:
        fmt = paragraph.paragraph_format

        if fmt.line_spacing:

            if isinstance(fmt.line_spacing, float):
                styles.append(
                    f"line-height:{fmt.line_spacing};"
                )

            elif hasattr(fmt.line_spacing, "twips"):
                px = twips_to_px(
                    fmt.line_spacing.twips
                )

                if px:
                    styles.append(
                        f"line-height:{px}px;"
                    )
    except Exception:
        pass

    # --------------------------------------------------------
    # KEEP TOGETHER / KEEP WITH NEXT
    # --------------------------------------------------------

    try:
        pPr = paragraph._p.get_or_add_pPr()

        keep_next = pPr.find(qn("w:keepNext"))

        if keep_next is not None:
            styles.append(
                "break-after:avoid;"
            )

    except Exception:
        pass

    return "".join(styles)


# ============================================================
# PAGE INFORMATION
# ============================================================

def get_page_css(document):
    """
    Citeste dimensiunea paginii si marginile din sectiunea Word.
    """

    try:
        section = document.sections[0]

        page_width = twips_to_px(
            section.page_width.twips
        )

        page_height = twips_to_px(
            section.page_height.twips
        )

        margin_top = twips_to_px(
            section.top_margin.twips
        )

        margin_bottom = twips_to_px(
            section.bottom_margin.twips
        )

        margin_left = twips_to_px(
            section.left_margin.twips
        )

        margin_right = twips_to_px(
            section.right_margin.twips
        )

        return {
            "width": page_width,
            "height": page_height,
            "margin_top": margin_top,
            "margin_bottom": margin_bottom,
            "margin_left": margin_left,
            "margin_right": margin_right
        }

    except Exception:
        return {
            "width": 794,
            "height": 1123,
            "margin_top": 96,
            "margin_bottom": 96,
            "margin_left": 96,
            "margin_right": 96
        }


# ============================================================
# IMAGE EXTRACTION
# ============================================================

def extract_document_images(document, prefix):
    """
    Extrage imaginile din DOCX si le salveaza separat.

    Returneaza:
        relationship_id -> URL
    """

    image_map = {}

    try:
        relationships = document.part.rels

        for rel_id, rel in relationships.items():

            if "image" not in rel.reltype:
                continue

            target = rel.target_part

            blob = target.blob

            filename = (
                f"{prefix}_{uuid.uuid4().hex[:12]}.png"
            )

            output_path = WORD_MEDIA_DIR / filename

            with open(output_path, "wb") as f:
                f.write(blob)

            image_map[rel_id] = (
                f"/files/media/{filename}"
            )

    except Exception:
        pass

    return image_map


# ============================================================
# DRAWING / IMAGE
# ============================================================

def drawing_to_html(run, image_map):
    """
    Detecteaza imaginea aflata intr-un run Word.
    """

    try:
        drawing = run._r.find(qn("w:drawing"))

        if drawing is None:
            return ""

        blips = drawing.findall(
            ".//" + qn("a:blip")
        )

        if not blips:
            return ""

        for blip in blips:

            rel_id = blip.get(
                qn("r:embed")
            )

            if rel_id and rel_id in image_map:

                src = image_map[rel_id]

                # ------------------------------------------------
                # INCERCAM SA AFLAM DIMENSIUNILE
                # ------------------------------------------------

                width = None
                height = None

                try:
                    extent = drawing.find(
                        ".//" + qn("wp:extent")
                    )

                    if extent is not None:

                        cx = extent.get("cx")
                        cy = extent.get("cy")

                        if cx:
                            width = round(
                                int(cx) / 9525
                            )

                        if cy:
                            height = round(
                                int(cy) / 9525
                            )

                except Exception:
                    pass

                attrs = [
                    f'src="{escape_html(src)}"',
                    'class="word-image"'
                ]

                if width:
                    attrs.append(
                        f'width="{width}"'
                    )

                if height:
                    attrs.append(
                        f'height="{height}"'
                    )

                return (
                    f'<img {" ".join(attrs)} '
                    f'alt="Imagine document Word">'
                )

    except Exception:
        pass

    return ""


# ============================================================
# RUN CONTENT
# ============================================================

def run_content_to_html(run, image_map):

    image_html = drawing_to_html(
        run,
        image_map
    )

    if image_html:
        return image_html

    return run_to_html(run)


# ============================================================
# PARAGRAPH TO HTML
# ============================================================

def paragraph_to_html(
    paragraph,
    paragraph_index,
    image_map
):
    """
    Important:

    NU convertim Heading 1/2/3 etc. in taguri XML.

    Pastram documentul asa cum este el in Word.

    Stilul Word este pastrat doar ca informatie
    si pentru randarea vizuala.
    """

    p_style = ""

    try:
        if paragraph.style:
            p_style = paragraph.style.name or ""
    except Exception:
        p_style = ""

    style_attr = get_paragraph_css(
        paragraph
    )

    # --------------------------------------------------------
    # CLASA VIZUALA
    # --------------------------------------------------------

    classes = [
        "word-paragraph"
    ]

    if p_style:
        safe_style = re.sub(
            r"[^a-zA-Z0-9_-]+",
            "-",
            p_style
        ).lower()

        classes.append(
            f"style-{safe_style}"
        )

    # --------------------------------------------------------
    # WORD STYLE
    # --------------------------------------------------------

    data_style = escape_html(
        p_style
    )

    # --------------------------------------------------------
    # CONTENT
    # --------------------------------------------------------

    content = ""

    for run in paragraph.runs:
        content += run_content_to_html(
            run,
            image_map
        )

    # --------------------------------------------------------
    # TABURI
    # --------------------------------------------------------

    content = content.replace(
        "\t",
        '<span class="word-tab"></span>'
    )

    # --------------------------------------------------------
    # EMPTY PARAGRAPH
    # --------------------------------------------------------

    if not content:
        content = "&nbsp;"

    return (
        f'<p '
        f'class="{" ".join(classes)}" '
        f'data-word-index="{paragraph_index}" '
        f'data-word-style="{data_style}" '
        f'style="{style_attr}">'
        f'{content}'
        f'</p>'
    )


# ============================================================
# TABLE CELL
# ============================================================

def table_cell_to_html(cell, image_map):

    content = ""

    for paragraph_index, paragraph in enumerate(
        cell.paragraphs
    ):
        content += paragraph_to_html(
            paragraph,
            paragraph_index,
            image_map
        )

    return (
        '<td class="word-table-cell">'
        f'{content}'
        '</td>'
    )


# ============================================================
# TABLE
# ============================================================

def table_to_html(table, table_index, image_map):

    rows_html = ""

    for row_index, row in enumerate(table.rows):

        cells_html = ""

        for cell in row.cells:

            cell_html = ""

            for paragraph_index, paragraph in enumerate(
                cell.paragraphs
            ):
                cell_html += paragraph_to_html(
                    paragraph,
                    paragraph_index,
                    image_map
                )

            tag = "th" if row_index == 0 else "td"

            cells_html += (
                f'<{tag} '
                f'class="word-table-cell">'
                f'{cell_html}'
                f'</{tag}>'
            )

        rows_html += (
            '<tr class="word-table-row">'
            f'{cells_html}'
            '</tr>'
        )

    return (
        '<div '
        'class="word-table-wrapper" '
        f'data-word-table="{table_index}">'
        '<table class="word-table">'
        f'{rows_html}'
        '</table>'
        '</div>'
    )


# ============================================================
# BLOCK ITEMS
# ============================================================

def iter_block_items(parent):
    """
    Pastreaza ordinea reala a documentului:

        paragraf
        paragraf
        tabel
        paragraf
        tabel
        etc.
    """

    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = parent.element.body

    for child in parent_elm.iterchildren():

        if child.tag == qn("w:p"):
            yield Paragraph(
                child,
                parent
            )

        elif child.tag == qn("w:tbl"):
            yield Table(
                child,
                parent
            )


# ============================================================
# DOCX -> HTML
# ============================================================

def docx_to_html(document, image_map):

    blocks = []

    paragraph_index = 0
    table_index = 0

    for block in iter_block_items(document):

        if block.__class__.__name__ == "Paragraph":

            blocks.append(
                paragraph_to_html(
                    block,
                    paragraph_index,
                    image_map
                )
            )

            paragraph_index += 1

        else:

            blocks.append(
                table_to_html(
                    block,
                    table_index,
                    image_map
                )
            )

            table_index += 1

    return "\n".join(blocks)


# ============================================================
# HOME
# ============================================================

@app.get(
    "/",
    response_class=HTMLResponse
)
async def home(request: Request):

    return templates.TemplateResponse(
        "index.html",
        {
            "request": request
        }
    )


# ============================================================
# UPLOAD WORD
# ============================================================

@app.post("/upload-word")
async def upload_word(
    file: UploadFile = File(...)
):

    if not file.filename:
        return JSONResponse(
            {
                "status": "error",
                "message": "Nu a fost selectat niciun fisier."
            },
            status_code=400
        )

    original_name = file.filename

    if not original_name.lower().endswith(".docx"):
        return JSONResponse(
            {
                "status": "error",
                "message": "Acceptam doar fisiere .docx."
            },
            status_code=400
        )

    safe_name = safe_filename(
        original_name
    )

    unique_name = (
        f"{uuid.uuid4().hex[:12]}_{safe_name}"
    )

    output_path = (
        WORD_UPLOADS_DIR / unique_name
    )

    try:

        # ----------------------------------------------------
        # SALVARE DOCX
        # ----------------------------------------------------

        with open(
            output_path,
            "wb"
        ) as buffer:

            shutil.copyfileobj(
                file.file,
                buffer
            )

        # ----------------------------------------------------
        # DESCHIDERE DOCUMENT
        # ----------------------------------------------------

        document = Document(
            str(output_path)
        )

        # ----------------------------------------------------
        # EXTRAGEM IMAGINILE
        # ----------------------------------------------------

        image_map = extract_document_images(
            document,
            Path(unique_name).stem
        )

        # ----------------------------------------------------
        # CONVERTIM IN HTML VIZUAL
        # ----------------------------------------------------

        html = docx_to_html(
            document,
            image_map
        )

        # ----------------------------------------------------
        # PAGE INFO
        # ----------------------------------------------------

        page = get_page_css(
            document
        )

        return JSONResponse(
            {
                "status": "success",
                "filename": original_name,
                "stored_filename": unique_name,
                "url": f"/files/word/{unique_name}",
                "html": html,
                "page": page
            }
        )

    except Exception as e:

        # daca apare eroare, stergem fisierul partial

        try:
            if output_path.exists():
                output_path.unlink()
        except Exception:
            pass

        return JSONResponse(
            {
                "status": "error",
                "message": (
                    "Eroare la procesarea documentului Word: "
                    f"{str(e)}"
                )
            },
            status_code=500
        )


# ============================================================
# HEALTH
# ============================================================

@app.get("/health")
async def health():

    return {
        "status": "ok",
        "app": "XML Tagger",
        "version": "3.0.0"
    }
